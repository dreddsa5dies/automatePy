#! python

import docx

doc = docx.Document('test.docx')

doc.paragraphs[0].style = 'Normal'
doc.add_paragraph('До чего дошли мы?', 'Номерной 2')
# С кириллицой не работает (
# doc.add_heading('Заголовок 1', 0)
doc.save('demo.docx')