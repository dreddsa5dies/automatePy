#! python

import docx

doc = docx.Document('test.docx')

for i in range(len(doc.paragraphs)):
    print(doc.paragraphs[i].text)

for i in range(len(doc.paragraphs)):
    for k in range(len(doc.paragraphs[i].runs)):
        print(doc.paragraphs[i].runs[k].text)